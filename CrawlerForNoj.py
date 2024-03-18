# @Time: 2024/3/14 13:15
# @Author: AsakawaNagi
# @File: CrawlerForNoj.py
# @Software: PyCharm
# @Note： 哥因为懒得安装脚本复制写了这个东西，怎么说吧，活该懒人统治世界，这个版本是黑框框的

import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from io import BytesIO

# def remove_english(text):
#     pattern = r'([\u4e00-\u9fa5]+)[a-zA-Z]+'
#     result = re.sub(pattern, r'\1', text)
#     return result
#我这里本来是想实现让用户选择是否需要保留英文说明，因为太占格子了，结果没能实现

# 按网上说的创建一个session对象，尽管我到现在也不知道有啥用
session = requests.Session()

headers = {
    'User-Agent': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 Edg/122.0.0.0"
}

username = input("请输入账户名: ")
password = input("请输入密码: ")

# 提交登录表单的数据
login_data = {
    'username': username,
    'password': password
}

# 提交登录表单的数据,这里直接使用我的noj账号密码好了。不对，不告诉你！
# login_data = {
#     'username': '',
#     'password': ''
# }

# 发送post请求登录，然后发现它甚至没有做反爬，不需要加header也可以爬出来结果
login_response = session.post('http://10.12.13.248/cpbox/', data=login_data, headers=headers)

# 检查登录
after_login_response = session.get('http://10.12.13.248/cpbox/webfile.aspx', headers=headers)
if login_response.ok and after_login_response.ok:
    print('登录成功！')
    # while True:
    #     english_input = input("是否需要英文描述？需要请输入1，不需要请输入0: ")
    #     if english_input in ('0', '1'):
    #         break
    #     else:
    #         print("请输入正确的选项")

    while True:
        inout_input = (int)(input("是否需要输入输出范例？需要请输入1，不需要请输入0: "))
        if (inout_input == 1) or (inout_input == 0):
            break
        else:
            print("请输入正确的选项")

    # 检查是否登陆成功，跳转到正确的noj主页面
    response = session.get('http://10.12.13.248/cpbox/cpNPUOJ.aspx#', headers=headers)
    if response.ok:
        doc = Document()
        # 确认一切正常 创建文件
        # 我最开始的思路是连续点击作业空间再点击noj习题，本来都做好了…………然后翻了下源码发现单独noj习题的链接被备注进去了，直接输入居然可以跳转，页面也简单多。听我说谢谢你。
        print("成功跳转到 http://10.12.13.248/cpbox/cpNPUOJ.aspx# 页面")
        html = response.text
        soup = BeautifulSoup(html, 'html.parser')

        name = soup.find("span", id="lblSTRealName")
        if name:
            print(name.text.strip(),"同学你好！")

        # 所有疑似习题链接共通点是<a>标签 id="cpIPPSFReader"
        all_link = soup.find_all("a", id="cpIPPSFReader")

        count = 1
        # 是不是感觉很莫名其妙，这是为了给题目输出计算序号

        for link in all_link:
            onclick_value = link.get('onclick')
            urls = re.findall(r"url:'(.*?)'", onclick_value)
            if urls:
                # 我在浏览器里尝试的时候，发现需要添加前缀才能成功跳转
                full_url = 'http://10.12.13.248/cpbox/' + urls[0]
                # noj你甚至不支持https协议
                # print(full_url)
                final_response = session.get(full_url, headers=headers)
                if final_response.ok:
                    soup = BeautifulSoup(final_response.text, 'html.parser')

                    # 获取题目名称
                    title = soup.find("span", id="lblIPPDFtitle")
                    if title:
                        print("题目", count, ":", title.text.strip())
                        doc.add_heading("题目" + str(count) + ": " + title.text.strip(), 2)
                        count += 1

                    # 获取描述
                    description = soup.find("span", id="lblIPPDFdescription")
                    if description:
                        print("描述:", description.text.strip())
                        doc.add_paragraph("描述: " + description.text.strip())

                    # 获取图片（如果有的话）
                    image = soup.find("img")
                    if image:
                        src = image.get("src")
                        if src:
                            full_match = 'http://10.12.13.248/cpbox/' + src
                            print(full_match)
                            try:
                                # 下载图像
                                response = requests.get(full_match)
                                image_bytes = BytesIO(response.content)

                                # 插入图像
                                doc.add_picture(image_bytes)
                                doc.add_paragraph()  # 添加一个空行分隔图片和其他内容
                            except Exception as e:
                                print("插入错误:", e)

                    # 获取输入
                    input_section = soup.find("span", id="lblIPPDFiutput")
                    if input_section:
                        # if english_input == '0':
                        #     # 如果不需要英语描述，就需要进行字符串切分
                        #     first = input_section.text.strip()
                        #     second = remove_english(first)
                        #     print("输入:", second)
                        # else:
                        #     print("输入:", input_section.text.strip())
                        # 笑死 写不对了 明知山有虎 猛敲退堂鼓
                        print("输入:", input_section.text.strip())
                        doc.add_paragraph("输入: " + input_section.text.strip())

                    # 获取输出
                    output_section = soup.find("span", id="lblIPPDFoutput")
                    if output_section:
                        print("输出:", output_section.text.strip())
                        doc.add_paragraph("输出: " + output_section.text.strip())

                    if inout_input:
                        # 获取样例输入
                        sampleintput_section = soup.find("span", id="lblIPPDFsampleinput")
                        if sampleintput_section:
                            print("样例输入:", sampleintput_section.text.strip())
                            doc.add_paragraph("样例输入: " + sampleintput_section.text.strip())

                        # 获取样例输出
                        sampleoutput_section = soup.find("span", id="lblIPPDFsampleoutput")
                        if sampleoutput_section:
                            print("样例输出:", sampleoutput_section.text.strip())
                            doc.add_paragraph("样例输出: " + sampleoutput_section.text.strip())

                    print("\n")
                    doc.add_paragraph("\n")

            else:
                print("未找到URL")
        doc.add_paragraph("\ncode by 外国语学院 2023303822 丁中惠")
        doc.save("当前noj题目.docx")

    else:
        print("跳转失败")

else:
    print('登录失败，状态码：', login_response.status_code)


# /html/body/div/div[2]/div[1]/div[1]/p[3]/span
#//*[@id="lblIPPDFdescription"]/text()[1]
# http://10.12.13.248/cpbox/cpNPUOJ.aspx#
#document.querySelector("#cpIPPSFReader")
# <span id="lblIPPDFdescription">输入两个整数A、B，输出它们的和。<br>Input two integers A and B, and output their sum.</span>
# {url:'cpIPPDFReader.aspx?v=%55%36%77%55%38%51%57%6f%43%65%6c%65%48%4c%34%64%35%54%2f%45%6b%66%36%4b%2b%4d%5a%71%68%6d%6a%4a%33%4e%5a%4c%49%2b%59%7a%50%34%59%3d'
#http://10.12.13.248/cpbox/cpIPPDFReader.aspx?v=%2f%73%58%54%4a%75%37%36%74%35%2f%46%7a%58%44%43%33%50%77%74%7a%79%35%52%71%43%4d%7a%7a%74%74%64%6b%4e%54%38%70%59%52%61%64%77%41%3d
