import tkinter as tk
from tkinter import messagebox
import requests
from bs4 import BeautifulSoup
from docx import Document
import re
from io import BytesIO

headers = {
    'User-Agent': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 Edg/122.0.0.0"
}


root = tk.Tk()
root.title("CrawlerForNoj")

space = tk.Label(root)
space.grid(row=0, column=0, sticky="e")

tip_label = tk.Label(root, text="""众所周知，noj最讨人厌的就是不能复制输入输出，尤其会在写字符串的时候导致频繁破防\n因此我写了这个程序，输入你的noj账号密码，你就可以获取所有noj题目，并且自定义是否需要输入输出\n它不会收集你的个人数据
 """, padx=10, wraplength=300, justify=tk.LEFT)
tip_label.grid(row=1, columnspan=2)

username_label = tk.Label(root, text="账户名:")
username_label.grid(row=2, column=0, sticky="e")
username_entry = tk.Entry(root)
username_entry.grid(row=2, column=1)

password_label = tk.Label(root, text="密码:")
password_label.grid(row=3, column=0, sticky="e")
password_entry = tk.Entry(root, show="*")
password_entry.grid(row=3, column=1)

space = tk.Label(root)
space.grid(row=4, column=0, sticky="e")

def login():
    username = username_entry.get()
    password = password_entry.get()

    session = requests.Session()

    login_data = {
        'username': username,
        'password': password
    }

    login_response = session.post('http://10.12.13.248/cpbox/', data=login_data, headers=headers)
    after_login_response = session.get('http://10.12.13.248/cpbox/webfile.aspx', headers=headers)
    if login_response.ok and after_login_response.ok:
        messagebox.showinfo("CrawlerForNoj", "登录成功！")
        # 开始爬了
        scrape(session)
    else:
        messagebox.showerror("CrawlerForNoj", f"登录失败，状态码: {login_response.status_code}")

def scrape(session):
    if messagebox.askyesno("CrawlerForNoj", "是否需要样例输入输出?"):
        sample_needed = True
    else:
        sample_needed = False

    response = session.get('http://10.12.13.248/cpbox/cpNPUOJ.aspx#')

    if response.ok:
        doc = Document()
        print("成功跳转到 http://10.12.13.248/cpbox/cpNPUOJ.aspx# 页面")
        html = response.text

        soup = BeautifulSoup(html, 'html.parser')
        name = soup.find("span", id="lblSTRealName")
        if name:
            messagebox.showinfo("Welcome", f"{name.text.strip()} 同学，欢迎使用！")

        soup = BeautifulSoup(html, 'html.parser')
        all_link = soup.find_all("a", id="cpIPPSFReader")

        count = 1

        for link in all_link:
            onclick_value = link.get('onclick')
            urls = re.findall(r"url:'(.*?)'", onclick_value)
            if urls:
                full_url = 'http://10.12.13.248/cpbox/' + urls[0]
                final_response = session.get(full_url, headers=headers)
                if final_response.ok:
                    soup = BeautifulSoup(final_response.text, 'html.parser')

                    title = soup.find("span", id="lblIPPDFtitle")
                    if title:
                        print("题目", count, ":", title.text.strip())
                        doc.add_heading("题目" + str(count) + ": " + title.text.strip(), 2)
                        count += 1

                    description = soup.find("span", id="lblIPPDFdescription")
                    if description:
                        print("描述:", description.text.strip())
                        doc.add_paragraph("描述: " + description.text.strip())

                    # 获取图片（如果有的话）
                    image = soup.find("img")
                    if image:
                        src = image.get("src")
                        if src:
                            full_match = 'http://10.12.13.248/cpbox/' + src
                            print(full_match)
                            try:
                                response = requests.get(full_match)
                                image_bytes = BytesIO(response.content)

                                doc.add_picture(image_bytes)
                            except Exception as e:
                                print("插入错误:", e)

                    input_section = soup.find("span", id="lblIPPDFiutput")
                    if input_section:
                        print("输入:", input_section.text.strip())
                        doc.add_paragraph("输入: " + input_section.text.strip())

                    output_section = soup.find("span", id="lblIPPDFoutput")
                    if output_section:
                        print("输出:", output_section.text.strip())
                        doc.add_paragraph("输出: " + output_section.text.strip())

                    if sample_needed:
                        sample_input_section = soup.find("span", id="lblIPPDFsampleinput")
                        if sample_input_section:
                            print("样例输入:", sample_input_section.text.strip())
                            doc.add_paragraph("样例输入: " + sample_input_section.text.strip())

                        sample_output_section = soup.find("span", id="lblIPPDFsampleoutput")
                        if sample_output_section:
                            print("样例输出:", sample_output_section.text.strip())
                            doc.add_paragraph("样例输出: " + sample_output_section.text.strip())

                    print("\n")
                    doc.add_paragraph("\n")

        doc.add_paragraph("\ncode by AsakawaNagi")
        doc.save("当前noj题目.docx")
        messagebox.showinfo("CrawlerForNoj", "生成成功，关闭全部对话框即可看到docx文件")

    else:
        print("跳转失败")


login_button = tk.Button(root, text="登录", command=login)
login_button.grid(row=5, columnspan=2)

space = tk.Label(root)
space.grid(row=6, column=0, sticky="e")

root.mainloop()
